import unittest
import pandas as pd
import os
from io import BytesIO
from code_exec import save_file, run_code, check_code, docx2tabular # Asegúrate de cambiar 'your_module' al nombre real del archivo donde está tu código.

class TestFileProcessing(unittest.TestCase):

    def setUp(self):
        # Crear un directorio para los archivos subidos si no existe
        self.uploaded_files_dir = 'uploaded_files'
        os.makedirs(self.uploaded_files_dir, exist_ok=True)

        # Crear archivos de prueba
        self.csv_file = f'{self.uploaded_files_dir}/test.csv'
        self.xlsx_file = f'{self.uploaded_files_dir}/test.xlsx'
        self.docx_file = f'{self.uploaded_files_dir}/test.docx'

        # Crear un archivo CSV de prueba
        df = pd.DataFrame({'A': [1, 2], 'B': [3, 4]})
        df.to_csv(self.csv_file, index=False)

        # Crear un archivo Excel de prueba
        df.to_excel(self.xlsx_file, index=False)

        # Crear un archivo DOCX de prueba
        from docx import Document
        doc = Document()
        doc.add_table(rows=2, cols=2)
        doc.tables[0].cell(0, 0).text = "Header1"
        doc.tables[0].cell(0, 1).text = "Header2"
        doc.tables[0].cell(1, 0).text = "Value1"
        doc.tables[0].cell(1, 1).text = "Value2"
        doc.save(self.docx_file)

    def tearDown(self):
        # Eliminar los archivos creados durante la prueba
        for filename in [self.csv_file, self.xlsx_file, self.docx_file]:
            if os.path.exists(filename):
                os.remove(filename)
        
    def test_docx2tabular(self):
        tabular_data = docx2tabular(self.docx_file)
        self.assertEqual(len(tabular_data), 2)  # Debe haber 2 filas
        self.assertEqual(tabular_data[0], ["Header1", "Header2"])  # Verifica encabezados

    def test_save_csv(self):
        with open(self.csv_file, 'rb') as f:
            uploadFile = BytesIO(f.read())
            uploadFile.name = 'test.csv'
            uploadFile.file_id = '12345'
            df, query_tabular, file_detail = save_file(uploadFile)
        
            self.assertIsInstance(df, pd.DataFrame)
            self.assertEqual(df.shape[0], 2)  # Verifica que hay 2 filas en el DataFrame
            self.assertIn('A', df.columns)     # Verifica que la columna 'A' está presente

    def test_save_xlsx(self):
        with open(self.xlsx_file, 'rb') as f:
            uploadFile = BytesIO(f.read())
            uploadFile.name = 'test.xlsx'
            uploadFile.file_id = '12345'
            df, query_tabular, file_detail = save_file(uploadFile)
        
            self.assertIsInstance(df, pd.DataFrame)
            self.assertEqual(df.shape[0], 2)  # Verifica que hay 2 filas en el DataFrame

    def test_save_docx(self):
        with open(self.docx_file, 'rb') as f:
            uploadFile = BytesIO(f.read())
            uploadFile.name = 'test.docx'
            uploadFile.file_id = '12345'
            df, query_tabular, file_detail = save_file(uploadFile)
        
            self.assertIsInstance(df, pd.DataFrame)
            self.assertEqual(df.shape[0], 1)  # Verifica que hay una fila en el DataFrame (encabezados)

    def test_run_code(self):
        code_to_run = """
a = [1, 2, 3]
print(a)
"""
        result = run_code(code_to_run, ['data.csv'], is_merge=False)
        self.assertEqual(result, [1, 2, 3])  # Verifica que el resultado es correcto

    def test_check_code(self):
        code_to_check = """
b = [4, 5]
print(b)
"""
        result = check_code(code_to_check, ['data.csv'], is_merge=False)
        self.assertTrue(result)  # Verifica que el código no arroja errores


if __name__ == '__main__':
    unittest.main()